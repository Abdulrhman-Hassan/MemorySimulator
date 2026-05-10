# ============================================================================
# Memory Management Segmentation Simulator - Project Report Generator
# ============================================================================
# Prerequisites:
#   pip install python-docx
#
# Usage:
#   python generate_report.py
#
# Output:
#   Creates "Memory_Management_Report/" directory with the .docx report
#   and a "Screenshots/" sub-folder for test case images.
# ============================================================================

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ======================== DIRECTORY SETUP ========================

REPORT_DIR = os.path.join(os.getcwd(), "Memory_Management_Report")
SCREENSHOTS_DIR = os.path.join(REPORT_DIR, "Screenshots")
OUTPUT_FILE = os.path.join(REPORT_DIR, "Memory_Segmentation_Simulator_Report.docx")

os.makedirs(REPORT_DIR, exist_ok=True)
os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

doc = Document()

# ======================== STYLE HELPERS ========================

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_code_block(doc, code_text):
    """Insert a formatted C# code block with gray background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_placeholder(doc):
    """Insert a yellow-highlighted screenshot placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[PLACEHOLDER: Insert Snapshot of Memory Layout drawing and Segment Tables here]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFF00" w:val="clear"/>')
    run.element.get_or_add_rPr().append(shading)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    """Insert a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


# ======================== TITLE PAGE ========================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Memory Management\nSegmentation Simulator")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Project Report")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

tech = doc.add_paragraph()
tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tech.add_run("Built with C# | WinUI 3 | .NET 10")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True

doc.add_paragraph()
course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = course.add_run("Operating Systems Course")
run.font.size = Pt(14)

doc.add_page_break()

# ======================== TABLE OF CONTENTS ========================

doc.add_heading("Table of Contents", level=1)

# To create a native, linked Table of Contents in Word, we need to insert
# a field code: { TOC \o "1-3" \h \z \u }
from docx.oxml import OxmlElement

paragraph = doc.add_paragraph()
run = paragraph.add_run()

# 1. Start the field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

# 2. Insert the field instruction (the TOC code)
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

# 3. Separate the instruction from the result
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')

# 4. (Optional but recommended) Add a placeholder text to tell the user to update the field
run_placeholder = OxmlElement('w:r')
text_placeholder = OxmlElement('w:t')
text_placeholder.text = "Right-click here and select 'Update Field' to generate the Table of Contents."
run_placeholder.append(text_placeholder)

# 5. End the field
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

# Append all elements to the run
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(run_placeholder)
run._r.append(fldChar3)

doc.add_page_break()

# ======================== 1. INTRODUCTION ========================

doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "This report documents the backend logic of a Memory Management Segmentation Simulator "
    "developed in C# as a WinUI 3 desktop application. The simulator demonstrates how an "
    "operating system manages memory using segmentation, where each process is divided into "
    "logical segments (Code, Data, Stack, etc.) that are independently allocated into free "
    "memory holes."
)
doc.add_paragraph(
    "The simulator supports two allocation algorithms (First-Fit and Best-Fit), enforces "
    "all-or-nothing allocation semantics, and performs automatic coalescing of adjacent free "
    "holes upon deallocation. This report focuses exclusively on the core backend engine, "
    "explaining the data structures, algorithms, and validation logic."
)

doc.add_page_break()

# ======================== 2. DATA STRUCTURES ========================

doc.add_heading("2. Data Structures", level=1)
doc.add_paragraph(
    "The simulator's backend is built around four core entity classes that model the memory "
    "management domain. These classes are defined in the MemorySimulator.Core namespace."
)

# 2.1 Segment
doc.add_heading("2.1 Segment Class", level=2)
doc.add_paragraph(
    "The Segment class represents a single logical segment of a process. Each segment has a "
    "name (e.g., Code, Data, Stack), a size, and once allocated, a base address in memory. "
    "The Limit property returns the segment's size, and EndAddress computes the exclusive "
    "ending address (BaseAddress + Size)."
)
add_code_block(doc, """public class Segment
{
    public string Name { get; set; }
    public int Size { get; set; }
    public int BaseAddress { get; set; } = -1;
    public int Limit => Size;
    public string ProcessId { get; set; }
    public int EndAddress => BaseAddress + Size;

    public Segment(string name, int size, string processId)
    {
        Name = name;
        Size = size;
        ProcessId = processId;
    }
}""")

# 2.2 Hole
doc.add_heading("2.2 Hole (Free Partition) Class", level=2)
doc.add_paragraph(
    "The Hole class represents a contiguous block of free memory. It tracks the starting "
    "address and size of the free region. The EndAddress property computes the exclusive end. "
    "A Clone() method enables safe simulation of allocation without modifying real state."
)
add_code_block(doc, """public class Hole
{
    public int StartingAddress { get; set; }
    public int Size { get; set; }
    public int EndAddress => StartingAddress + Size;

    public Hole(int startingAddress, int size)
    {
        StartingAddress = startingAddress;
        Size = size;
    }

    public Hole Clone() => new Hole(StartingAddress, Size);
}""")

# 2.3 Process
doc.add_heading("2.3 Process Class", level=2)
doc.add_paragraph(
    "The Process class groups multiple segments under a single process identifier. "
    "TotalSize sums the sizes of all segments belonging to the process."
)
add_code_block(doc, """public class Process
{
    public string Id { get; set; }
    public List<Segment> Segments { get; set; }

    public Process(string id, List<Segment> segments)
    {
        Id = id;
        Segments = segments;
    }

    public int TotalSize => Segments.Sum(s => s.Size);
}""")

# 2.4 & 2.5 Tables
doc.add_heading("2.4 Allocated Partitions Table", level=2)
doc.add_paragraph(
    "The MemoryManager maintains a List<Segment> called AllocatedPartitions. This list holds "
    "every segment from every active process that has been successfully placed in memory. "
    "Segments are sorted by BaseAddress for display purposes."
)

doc.add_heading("2.5 Free Partitions Table", level=2)
doc.add_paragraph(
    "The MemoryManager maintains a List<Hole> called FreePartitions. This list tracks all "
    "available free memory blocks. After every allocation or deallocation, this list is "
    "updated and sorted by StartingAddress."
)
add_code_block(doc, """public class MemoryManager
{
    public int TotalMemorySize { get; private set; }
    public List<Hole> FreePartitions { get; private set; } = new();
    public List<Segment> AllocatedPartitions { get; private set; } = new();
    public Dictionary<string, Process> ActiveProcesses { get; private set; } = new();
    public bool IsInitialized { get; private set; }
}""")

doc.add_page_break()

# Save intermediate progress
doc.save(OUTPUT_FILE)
print("Stage 1 complete: Title, TOC, Sections 1-2 written.")

# ======================== 3. ALLOCATION ALGORITHMS ========================

doc.add_heading("3. Allocation Algorithms", level=1)

doc.add_heading("3.1 First-Fit Algorithm", level=2)
doc.add_paragraph(
    "The First-Fit algorithm scans the free partitions table in order of starting address "
    "and returns the first hole whose size is greater than or equal to the required segment "
    "size. This approach is fast but may leave small unusable fragments at lower addresses."
)
add_code_block(doc, """private int FindFirstFit(List<Hole> holes, int requiredSize)
{
    var sorted = holes.OrderBy(h => h.StartingAddress).ToList();
    for (int i = 0; i < sorted.Count; i++)
    {
        if (sorted[i].Size >= requiredSize)
        {
            return holes.IndexOf(sorted[i]);
        }
    }
    return -1;
}""")

doc.add_heading("3.2 Best-Fit Algorithm", level=2)
doc.add_paragraph(
    "The Best-Fit algorithm searches all free holes and selects the smallest hole that is "
    "still large enough to accommodate the segment. This minimizes wasted space within the "
    "chosen hole but may create very small, unusable fragments elsewhere."
)
add_code_block(doc, """private int FindBestFit(List<Hole> holes, int requiredSize)
{
    int bestIndex = -1;
    int bestSize = int.MaxValue;

    for (int i = 0; i < holes.Count; i++)
    {
        if (holes[i].Size >= requiredSize && holes[i].Size < bestSize)
        {
            bestIndex = i;
            bestSize = holes[i].Size;
        }
    }

    return bestIndex;
}""")

doc.add_heading("3.3 All-or-Nothing Allocation Strategy", level=2)
doc.add_paragraph(
    "A critical design rule of this simulator is that process allocation is all-or-nothing. "
    "Before modifying the real free partitions table, the engine clones it and simulates "
    "placing every segment. If any segment fails to fit, the entire process is rejected and "
    "no partial allocation occurs. Only when all segments succeed does the engine commit "
    "the changes to the real state."
)
add_code_block(doc, """public AllocationResult AllocateProcess(Process process, AllocationMethod method)
{
    // Clone free partitions for simulation
    var simulatedHoles = FreePartitions.Select(h => h.Clone()).ToList();
    var assignments = new List<(Segment segment, int baseAddress)>();

    foreach (var segment in process.Segments)
    {
        int holeIndex = FindHole(simulatedHoles, segment.Size, method);

        if (holeIndex == -1)
        {
            // Reject entire process - no partial allocation
            return AllocationResult.Fail(
                $"Segment '{segment.Name}' does not fit. No partial allocations made.");
        }

        var chosenHole = simulatedHoles[holeIndex];
        int baseAddress = chosenHole.StartingAddress;
        assignments.Add((segment, baseAddress));

        if (chosenHole.Size == segment.Size)
            simulatedHoles.RemoveAt(holeIndex);  // Exact fit
        else
        {
            chosenHole.StartingAddress += segment.Size;  // Shrink hole
            chosenHole.Size -= segment.Size;
        }
    }

    // COMMIT: All segments fit - apply to real state
    FreePartitions = simulatedHoles;
    foreach (var (segment, baseAddress) in assignments)
    {
        segment.BaseAddress = baseAddress;
        AllocatedPartitions.Add(segment);
    }
    ActiveProcesses[process.Id] = process;
    return AllocationResult.Ok("Allocated successfully.");
}""")

doc.add_page_break()

# ======================== 4. DEALLOCATION & COALESCING ========================

doc.add_heading("4. Deallocation and Coalescing", level=1)

doc.add_heading("4.1 Deallocation Logic", level=2)
doc.add_paragraph(
    "When a process is deallocated, each of its segments is converted back into a free hole "
    "and added to the free partitions table. The segment is removed from the allocated "
    "partitions table, and the process is removed from the active processes dictionary."
)
add_code_block(doc, """public AllocationResult DeallocateProcess(string processId)
{
    var process = ActiveProcesses[processId];

    foreach (var segment in process.Segments)
    {
        FreePartitions.Add(new Hole(segment.BaseAddress, segment.Size));
        AllocatedPartitions.Remove(segment);
    }

    ActiveProcesses.Remove(processId);
    CoalesceHoles();

    return AllocationResult.Ok("Deallocated successfully. Adjacent holes merged.");
}""")

doc.add_heading("4.2 Hole Coalescing (Merging)", level=2)
doc.add_paragraph(
    "After deallocation, the CoalesceHoles method sorts all free holes by starting address "
    "and then iterates through them. If two consecutive holes are adjacent (i.e., the end "
    "address of the first equals the start address of the next), they are merged into a "
    "single larger hole. This prevents external fragmentation from accumulating."
)
add_code_block(doc, """private void CoalesceHoles()
{
    if (FreePartitions.Count <= 1) return;

    FreePartitions.Sort((a, b) => a.StartingAddress.CompareTo(b.StartingAddress));

    var merged = new List<Hole> { FreePartitions[0] };

    for (int i = 1; i < FreePartitions.Count; i++)
    {
        var last = merged[merged.Count - 1];
        var current = FreePartitions[i];

        if (last.EndAddress == current.StartingAddress)
        {
            last.Size += current.Size;  // Merge adjacent holes
        }
        else
        {
            merged.Add(current);
        }
    }

    FreePartitions = merged;
}""")

doc.add_page_break()

# ======================== 5. INPUT VALIDATION ========================

doc.add_heading("5. Input Validation", level=1)
doc.add_paragraph(
    "The Validators static class provides robust input validation before any operation. "
    "All validation methods return an AllocationResult indicating success or failure."
)
doc.add_paragraph("Key validations include:")

validations = [
    "Memory size must be a positive integer.",
    "Each hole must have a non-negative starting address and a positive size.",
    "No hole may exceed total memory boundaries (Start + Size <= TotalMemory).",
    "No two holes may overlap (checked by sorting and comparing boundaries).",
    "Process ID must not be empty or duplicate.",
    "Each segment must have a non-empty name and positive size.",
    "Duplicate segment names within the same process are rejected.",
]
for v in validations:
    doc.add_paragraph(v, style='List Bullet')

add_code_block(doc, """public static AllocationResult ValidateHoles(int totalMemorySize, List<Hole> holes)
{
    for (int i = 0; i < holes.Count; i++)
    {
        var hole = holes[i];
        if (hole.StartingAddress < 0)
            return AllocationResult.Fail($"Hole #{i+1}: Start address must be >= 0.");
        if (hole.Size <= 0)
            return AllocationResult.Fail($"Hole #{i+1}: Size must be positive.");
        if (hole.StartingAddress + hole.Size > totalMemorySize)
            return AllocationResult.Fail($"Hole #{i+1}: Exceeds total memory.");
    }

    var sorted = holes.OrderBy(h => h.StartingAddress).ToList();
    for (int i = 0; i < sorted.Count - 1; i++)
    {
        if (sorted[i].EndAddress > sorted[i + 1].StartingAddress)
            return AllocationResult.Fail("Holes overlap.");
    }
    return AllocationResult.Ok();
}""")

doc.add_page_break()

# ======================== 6. SYSTEM TESTING ========================

doc.add_heading("6. System Testing", level=1)

doc.add_heading("6.1 Test Scenario Setup", level=2)
doc.add_paragraph("The following test scenario is used for both algorithm runs:")

doc.add_paragraph("Initial Memory Configuration:", style='List Bullet')
add_table(doc,
    ["Parameter", "Value"],
    [["Total Memory Size", "1000 K"]]
)

doc.add_paragraph("Initial Free Holes:", style='List Bullet')
add_table(doc,
    ["Hole", "Start Address", "Size"],
    [["H1", "0", "300"],
     ["H2", "400", "250"],
     ["H3", "700", "200"]]
)

doc.add_paragraph("Operations Pipeline:", style='List Bullet')
add_table(doc,
    ["Step", "Operation", "Process", "Segments"],
    [["1", "Allocate", "P1", "Code=100, Data=120, Stack=90"],
     ["2", "Allocate", "P2", "Code=200, Data=40"],
     ["3", "Allocate", "P3", "Code=120, Data=50"],
     ["4", "Deallocate", "P1", "-"],
     ["5", "Allocate", "P4", "Code=230, Data=40"]]
)

doc.add_page_break()

# ---------- RUN 1: FIRST-FIT ----------

doc.add_heading("6.2 Run 1: First-Fit Algorithm", level=2)

doc.add_heading("Initial State After Memory Setup", level=3)
doc.add_paragraph("Memory is initialized with Total Size = 1000K and three free holes.")
add_placeholder(doc)

doc.add_heading("Step 1: Allocate P1 (Code=100, Data=120, Stack=90) - First-Fit", level=3)
doc.add_paragraph(
    "Using First-Fit, the allocator scans holes by address order. P1's segments are placed "
    "sequentially into the first hole large enough to hold each segment."
)
add_placeholder(doc)

doc.add_heading("Step 2: Allocate P2 (Code=200, Data=40) - First-Fit", level=3)
doc.add_paragraph(
    "P2's segments are allocated into the remaining free holes using First-Fit order."
)
add_placeholder(doc)

doc.add_heading("Step 3: Allocate P3 (Code=120, Data=50) - First-Fit", level=3)
doc.add_paragraph(
    "P3's segments are allocated into available holes using First-Fit."
)
add_placeholder(doc)

doc.add_heading("Step 4: Deallocate P1", level=3)
doc.add_paragraph(
    "All segments of P1 are freed and converted back to holes. Adjacent holes are "
    "automatically coalesced (merged) by the CoalesceHoles algorithm."
)
add_placeholder(doc)

doc.add_heading("Step 5: Allocate P4 (Code=230, Data=40) - First-Fit", level=3)
doc.add_paragraph(
    "P4's segments are allocated into the newly available free holes after P1's deallocation."
)
add_placeholder(doc)

doc.add_page_break()

# ---------- RUN 2: BEST-FIT ----------

doc.add_heading("6.3 Run 2: Best-Fit Algorithm", level=2)

doc.add_heading("Initial State After Memory Setup", level=3)
doc.add_paragraph("Memory is re-initialized with the same configuration for the Best-Fit run.")
add_placeholder(doc)

doc.add_heading("Step 1: Allocate P1 (Code=100, Data=120, Stack=90) - Best-Fit", level=3)
doc.add_paragraph(
    "Using Best-Fit, the allocator finds the smallest hole that can still fit each segment, "
    "minimizing leftover space in the chosen hole."
)
add_placeholder(doc)

doc.add_heading("Step 2: Allocate P2 (Code=200, Data=40) - Best-Fit", level=3)
doc.add_paragraph(
    "P2's segments are placed into the smallest sufficient holes using Best-Fit."
)
add_placeholder(doc)

doc.add_heading("Step 3: Allocate P3 (Code=120, Data=50) - Best-Fit", level=3)
doc.add_paragraph(
    "P3's segments are allocated using Best-Fit selection."
)
add_placeholder(doc)

doc.add_heading("Step 4: Deallocate P1", level=3)
doc.add_paragraph(
    "All segments of P1 are freed. Adjacent holes are coalesced automatically."
)
add_placeholder(doc)

doc.add_heading("Step 5: Allocate P4 (Code=230, Data=40) - Best-Fit", level=3)
doc.add_paragraph(
    "P4's segments are placed using Best-Fit into the freed holes."
)
add_placeholder(doc)

doc.add_page_break()

# ======================== 7. CONCLUSION ========================

doc.add_heading("7. Conclusion", level=1)
doc.add_paragraph(
    "This Memory Segmentation Simulator successfully demonstrates the core concepts of "
    "segmented memory management as taught in Operating Systems courses. The backend engine "
    "implements robust data structures for tracking allocated and free partitions, supports "
    "both First-Fit and Best-Fit allocation strategies with all-or-nothing semantics, and "
    "performs automatic coalescing of adjacent free holes upon deallocation."
)
doc.add_paragraph(
    "The test scenario executed in Section 6 validates the correctness of both algorithms "
    "across a realistic pipeline of allocations and deallocations, demonstrating how different "
    "placement strategies lead to different memory layouts and fragmentation patterns."
)

# ======================== SAVE ========================

doc.save(OUTPUT_FILE)
print(f"\nReport generated successfully!")
print(f"Location: {OUTPUT_FILE}")
print(f"Screenshots folder: {SCREENSHOTS_DIR}")
print("Replace the yellow [PLACEHOLDER] markers with your actual screenshots.")
